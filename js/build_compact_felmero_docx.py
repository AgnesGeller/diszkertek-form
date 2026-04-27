from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import build_bovitett_felmero_docx as base


OUTPUT_DIR = Path(r"C:\Users\darli\bootcamp\FREEE\diszkertek-form\js\docs")
OUTPUT_PATH = OUTPUT_DIR / "form.docx"

GREEN = base.GREEN
GREEN_DARK = base.GREEN_DARK
SAND = base.SAND
TAUPE = base.TAUPE
GRAY = base.GRAY
BASE_SECTIONS = base.SECTIONS
EXTRA_SECTIONS = [
    {
        "title": "Vízelvezetés és csapadékvíz-kezelés",
        "metrics": [
            "Érintett felület (m2)",
            "Folyóka hossza (fm)",
            "Drénhossz (fm)",
            "Szikkasztó térfogat (m3)",
            "Lejtési irányok (db)",
            "Vízgyűjtő pontok (db)",
        ],
        "groups": [
            {
                "title": "Kiinduló probléma",
                "lines": [
                    "□ pangó víz □ sáros felület □ fal melletti visszafröccsenés □ lefolyási gond",
                    "□ pince / alap védelme fontos □ burkolat melletti vízterhelés",
                    "Csapadékvíz jelenlegi útvonala ismert: □ igen □ részben □ nem",
                    "Megjegyzés / kritikus pont: __________________________________________",
                ],
            },
            {
                "title": "Műszaki megoldások",
                "lines": [
                    "□ folyóka □ pontszerű összefolyó □ dréncső □ szikkasztó □ esővízgyűjtő",
                    "□ kavicsárok □ vízterelő sáv □ burkolatlejtés korrekció",
                    "Közmű vagy gyökér akadályozza: □ igen □ nem",
                    "Csatornába kötés lehetősége: □ igen □ nem",
                ],
            },
        ],
    },
    {
        "title": "Parkoló és gépkocsi-beálló",
        "metrics": [
            "Parkolóhelyek száma (db)",
            "Burkolt felület (m2)",
            "Szélesség (m)",
            "Hossz (m)",
            "Terhelési szint",
            "Fordulóív / manőverterület",
        ],
        "groups": [
            {
                "title": "Használat",
                "lines": [
                    "□ egy autó □ több autó □ vendégparkoló □ teherautó / utánfutó is",
                    "Fedett vagy nyitott kialakítás: □ fedett □ nyitott □ később fedhető",
                    "Tolatási / fordulási hely kritikus: □ igen □ nem",
                    "Téli használat kiemelt: □ igen □ nem",
                ],
            },
            {
                "title": "Kapcsolódó igények",
                "lines": [
                    "□ kapunyitás előkészítés □ kültéri áram □ világítás □ vízelvezetés",
                    "□ járda csatlakozás □ lépcső / rámpa □ kerékpártárolás",
                    "Meglévő beálló bontása: □ igen □ nem",
                    "Esztétikai szint: □ egyszerű □ rendezett □ reprezentatív",
                ],
            },
        ],
    },
    {
        "title": "Kültéri lépcső és rámpa",
        "metrics": [
            "Lépcsők száma (db)",
            "Szintkülönbség (cm)",
            "Szélesség (cm)",
            "Hossz (fm)",
            "Pihenők száma (db)",
            "Korlát szükséges-e",
        ],
        "groups": [
            {
                "title": "Használati szempontok",
                "lines": [
                    "□ fő közlekedési útvonal □ mellékútvonal □ teraszkapcsolat □ kertkapcsolat",
                    "□ idősek számára is kényelmes □ gyermekbarát □ akadálycsökkentett",
                    "Csúszásmentesség kiemelten fontos: □ igen □ nem",
                    "Éjszakai megvilágítás kell: □ igen □ nem",
                ],
            },
            {
                "title": "Kialakítás",
                "lines": [
                    "□ monolit □ burkolt lépcső □ terméskő □ térkő □ fa / WPC",
                    "□ rámpa □ kombinált lépcső-rámpa □ támfalhoz kapcsolódó",
                    "Korlát vagy oldalfal kell: □ igen □ nem",
                    "Vízelvezetés megoldandó: □ igen □ nem",
                ],
            },
        ],
    },
    {
        "title": "Tárolók és kiszolgáló elemek",
        "metrics": [
            "Tárolók száma (db)",
            "Alapterület (m2)",
            "Szélesség (m)",
            "Hossz (m)",
            "Magasság (m)",
            "Kiszolgáló funkciók (db)",
        ],
        "groups": [
            {
                "title": "Funkció",
                "lines": [
                    "□ szerszámtároló □ géptároló □ kerékpártároló □ kukatároló □ tűzifatároló",
                    "□ kerti bútor tárolás □ medencegépészet □ vegyes használat",
                    "Külön zárhatóság szükséges: □ igen □ nem",
                    "Külső megjelenés illeszkedjen a házhoz: □ igen □ nem",
                ],
            },
            {
                "title": "Kialakítás és infrastruktúra",
                "lines": [
                    "□ fa □ fém □ tégla □ könnyűszerkezet □ előregyártott",
                    "□ alapozás □ burkolat □ világítás □ áramkiállás □ vízkiállás",
                    "Szellőzés / gépészeti hozzáférés kell: □ igen □ nem",
                    "Helyigény / elhelyezési korlát: ______________________________________",
                ],
            },
        ],
    },
    {
        "title": "Veteményes és emelt ágyás",
        "metrics": [
            "Ágyások száma (db)",
            "Összterület (m2)",
            "Szélesség (m)",
            "Hossz (m)",
            "Magasság (cm)",
            "Öntözött felület (m2)",
        ],
        "groups": [
            {
                "title": "Használat",
                "lines": [
                    "□ veteményes □ fűszernövényes □ magaságyás □ dísz és haszon együtt",
                    "□ családi használat □ intenzív termesztés □ oktató / bemutató cél",
                    "Napos fekvés biztosított: □ igen □ részben □ nem",
                    "Szezonhosszabbítás igényelt: □ igen □ nem",
                ],
            },
            {
                "title": "Kialakítás",
                "lines": [
                    "□ fa □ fém □ tégla □ WPC □ kő",
                    "□ öntözés □ talajcsere □ komposztáló □ takarás □ kártevővédelem",
                    "Háló / fólia / árnyékolás kell: □ igen □ nem",
                    "Terménykezelési megjegyzés: _________________________________________",
                ],
            },
        ],
    },
    {
        "title": "Medence és jacuzzi előkészítés",
        "metrics": [
            "Egységek száma (db)",
            "Felület (m2)",
            "Gépészeti pontok (db)",
            "Áramigény",
            "Vízpontok (db)",
            "Kiszolgáló zónák (db)",
        ],
        "groups": [
            {
                "title": "Projektkör",
                "lines": [
                    "□ csak előkészítés □ teljes kivitelezés □ gépészeti fogadás □ burkolati kapcsolat",
                    "□ medence □ jacuzzi □ dézsa □ több funkció együtt",
                    "Későbbi bővíthetőség fontos: □ igen □ nem",
                    "Gyermekbiztonság kritikus: □ igen □ nem",
                ],
            },
            {
                "title": "Kapcsolódó igények",
                "lines": [
                    "□ gépészeti akna □ kültéri zuhany □ világítás □ árnyékolás □ pihenőterasz",
                    "□ vízelvezetés □ takarás □ tároló □ kültéri áram",
                    "Szervizhozzáférés biztosítandó: □ igen □ nem",
                    "Burkolati és esztétikai csatlakozás: __________________________________",
                ],
            },
        ],
    },
    {
        "title": "Kerti konyha, grillező és tűzrakó",
        "metrics": [
            "Egységek száma (db)",
            "Felület (m2)",
            "Pult hossza (fm)",
            "Vízpontok (db)",
            "Árampontok (db)",
            "Ülőhelyek száma (db)",
        ],
        "groups": [
            {
                "title": "Használat",
                "lines": [
                    "□ grillező □ komplett kerti konyha □ kemence □ tűzrakó □ bárpult",
                    "□ családi használat □ vendégfogadás □ rendszeres főzés",
                    "Füstirány és szomszédhatás érzékeny: □ igen □ nem",
                    "Időjárás elleni védelem kell: □ igen □ nem",
                ],
            },
            {
                "title": "Közmű és kiszolgálás",
                "lines": [
                    "□ víz □ csatorna □ áram □ világítás □ gáz / palackos gáz □ tárolás",
                    "□ hűtő □ mosogató □ fedett pult □ ülőfal / étkezőzóna",
                    "Meglévő teraszhoz kapcsolódik: □ igen □ nem",
                    "Karbantartási és takarítási elvárás: _________________________________",
                ],
            },
        ],
    },
    {
        "title": "Okoskert, automatika és biztonság",
        "metrics": [
            "Automatikák száma (db)",
            "Vezérlési körök (db)",
            "Kamerapontok (db)",
            "Szenzorok (db)",
            "Árampontok (db)",
            "Hálózati pontok (db)",
        ],
        "groups": [
            {
                "title": "Automatizálás",
                "lines": [
                    "□ robotfűnyíró □ automata kapu □ öntözésvezérlés □ világításvezérlés",
                    "□ időzítés □ távoli elérés □ telefonos alkalmazás □ központi vezérlés",
                    "Későbbi bővítésre előkészítés kell: □ igen □ nem",
                    "Hálózati / wifi lefedettség adott: □ igen □ részben □ nem",
                ],
            },
            {
                "title": "Biztonság és felügyelet",
                "lines": [
                    "□ kamera □ mozgásérzékelő □ kültéri riasztás □ kaputelefon □ világítási jelenet",
                    "□ gyermekbiztonság □ kutyabiztos kert □ lezárható zónák",
                    "Kritikus védendő terület: ____________________________________________",
                    "Jogosultsági / használói megjegyzés: _________________________________",
                ],
            },
        ],
    },
]
SECTIONS = BASE_SECTIONS + EXTRA_SECTIONS


def add_header(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    base.add_text(paragraph, "Helyszíni felmérő - Díszkertek", size=9, bold=True, color=GREEN_DARK)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "D9CDBB")
    border.append(bottom)
    paragraph._p.get_or_add_pPr().append(border)


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.1)
    section.right_margin = Cm(1.1)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.45)
    add_header(section)

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    style.font.size = Pt(8)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    base.add_text(p, "Helyszíni felmérő - Díszkertek", size=20, bold=True, color=GREEN_DARK)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(7)
    base.add_text(meta, "Bővített, papíralapú kertépítői helyszíni felmérő. Kézi kitöltésre, gyors terepi használatra.", size=8, color=TAUPE)

    info = doc.add_table(rows=1, cols=3)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    labels = ["Dátum", "Helyszín", "Kapcsolattartó"]
    for i, label in enumerate(labels):
        cell = info.cell(0, i)
        cell.width = Cm(8.6)
        base.set_cell_shading(cell, SAND)
        base.set_cell_border(cell, color="D8CCBA", size="7")
        base.set_cell_margins(cell, top=60, start=85, bottom=60, end=85)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(1)
        base.add_text(p1, label, size=8, bold=True, color=GREEN_DARK)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        base.add_text(p2, "__________________________", size=8.5, color=TAUPE)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_section_band(doc, title):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(25.9)
    base.set_cell_shading(cell, GREEN)
    base.set_cell_border(cell, color=GREEN, size="4")
    base.set_cell_margins(cell, top=45, start=95, bottom=45, end=95)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    base.add_text(p, title, size=11.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))


def add_compact_metrics(doc, metrics):
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(8.63)
            base.set_cell_shading(cell, "FBF9F5")
            base.set_cell_border(cell, color="DED3C2", size="7")
            base.set_cell_margins(cell, top=45, start=70, bottom=45, end=70)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for idx, metric in enumerate(metrics[:6]):
        cell = table.cell(idx // 3, idx % 3)
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0.5)
        base.add_text(p1, metric, size=7.2, bold=True, color=GREEN_DARK)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        base.add_text(p2, "__________________", size=8, color=TAUPE)


def add_solution_preference(doc, title="Megoldási szint"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(25.9)
    base.set_cell_shading(cell, "FBF9F5")
    base.set_cell_border(cell, color="DED3C2", size="6")
    base.set_cell_margins(cell, top=45, start=80, bottom=45, end=80)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    base.add_text(
        p,
        f"{title}: □ költséghatékony / megtakarításra optimalizált  □ kiegyensúlyozott  □ teljes műszaki tartalom / magasabb esztétikai szint",
        size=7.1,
        color=GRAY,
    )


def add_compact_groups(doc, groups, is_alapadatok=False):
    cols = 2
    rows = (len(groups) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    width = Cm(12.95)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.width = width
            base.set_cell_shading(cell, "FFFFFF")
            base.set_cell_border(cell, color="E3D8C8", size="6")
            base.set_cell_margins(cell, top=50, start=75, bottom=50, end=75)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            idx = i * cols + j
            if idx >= len(groups):
                cell.text = ""
                continue
            group = groups[idx]
            head = cell.paragraphs[0]
            head.paragraph_format.space_after = Pt(1)
            base.add_text(head, group["title"], size=7.9, bold=True, color=GREEN_DARK)
            limit = 5 if is_alapadatok else 4
            for line in group["lines"][:limit]:
                p = cell.add_paragraph()
                p.paragraph_format.space_after = Pt(0.1)
                p.paragraph_format.line_spacing = 0.95
                base.add_text(p, base.paper_line(line), size=6.7 if is_alapadatok else 6.9, color=GRAY)


def add_compact_notes(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0.5)
    base.add_text(p, "Megjegyzés:", size=7.8, bold=True, color=GREEN_DARK)
    for _ in range(2):
        line = doc.add_paragraph()
        line.paragraph_format.space_before = Pt(0)
        line.paragraph_format.space_after = Pt(0.5)
        base.add_text(line, "__________________________________________________________________________________________", size=7.5, color=TAUPE)


def add_section(doc, section_data):
    add_section_band(doc, section_data["title"])
    add_compact_metrics(doc, section_data["metrics"])
    if section_data["title"] != "Alapadatok":
        add_solution_preference(doc)
    add_compact_groups(doc, section_data["groups"], is_alapadatok=section_data["title"] == "Alapadatok")
    add_compact_notes(doc)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Helyszíni felmérő - Díszkertek"
    doc.core_properties.subject = "Bővített papíralapú kertépítői felmérő"
    doc.core_properties.author = "Codex"

    add_cover(doc)
    for section in SECTIONS:
        add_section(doc, section)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
