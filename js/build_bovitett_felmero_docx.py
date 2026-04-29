import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent / "docs"
OUTPUT_PATH = OUTPUT_DIR / "Diszkertek_Bovitett_Helyszini_Felmero_Papir.docx"

GREEN = "335C4E"
GREEN_DARK = RGBColor(0x23, 0x42, 0x37)
SAND = "F5F1E8"
SAND_DARK = "E4D8C4"
TAUPE = RGBColor(0x73, 0x63, 0x52)
GRAY = RGBColor(0x63, 0x63, 0x63)

HU_REPLACEMENTS = [
    ("DISZKERTEK | BOVITETT KERTEPITOI HELYSZINI FELMERO", "DÍSZKERTEK | BŐVÍTETT KERTÉPÍTŐI HELYSZÍNI FELMÉRŐ"),
    ("Diszkertek bovitett helyszini felmero", "Díszkertek bővített helyszíni felmérő"),
    ("Kertepitoi helyszini felmero lap", "Kertépítői helyszíni felmérő lap"),
    ("Nyomtathato, pipalhato, rovid megjegyzesekkel kitoltheto oldal.", "Nyomtatható, pipálható, rövid megjegyzésekkel kitölthető oldal."),
    ("Megjegyzes / helyszini eszrevetel", "Megjegyzés / helyszíni észrevétel"),
    ("Alapadatok", "Alapadatok"),
    ("Ontozorendszer", "Öntözőrendszer"),
    ("Burkolatok", "Burkolatok"),
    ("Pergolak", "Pergolák"),
    ("Jatszoteri elemek", "Játszótéri elemek"),
    ("Agyasok", "Ágyások"),
    ("Sovenyek", "Sövények"),
    ("Fa ultetes", "Fa ültetés"),
    ("Szegely epites", "Szegély építés"),
    ("Fuvesites", "Füvesítés"),
    ("Bontas", "Bontás"),
    ("Tamfal epites", "Támfal építés"),
    ("Kerti vilagitas", "Kerti világítás"),
    ("Kerti aramforras", "Kerti áramforrás"),
    ("Telektisztitas", "Telektisztítás"),
    ("Tereprendezes", "Tereprendezés"),
    ("Sziklakert epites", "Sziklakert építés"),
    ("Diszkavics agyas epites", "Díszkavics ágyás építés"),
    ("Kerites epites", "Kerítés építés"),
    ("Kerti to epites", "Kerti tó építés"),
    ("Tervrajz keszites", "Tervrajz készítés"),
    ("Latvanyterv keszites", "Látványterv készítés"),
    ("Kapcsolat es helyszin", "Kapcsolat és helyszín"),
    ("Megkozelites es kozmuvek", "Megközelítés és közművek"),
    ("Atadott dokumentumok", "Átadott dokumentumok"),
    ("Hasznalati igenyek", "Használati igények"),
    ("Stilus es terkapcsolat", "Stílus és térkapcsolat"),
    ("Fenntartas es uzemeltetes", "Fenntartás és üzemeltetés"),
    ("Koltsegkeret es utemezes", "Költségkeret és ütemezés"),
    ("Projektosszefoglalo", "Projektösszefoglaló"),
    ("Rendszer jellege", "Rendszer jellege"),
    ("Ontozesi mod", "Öntözési mód"),
    ("Kivitelezesi korulmenyek", "Kivitelezési körülmények"),
    ("Burkolat anyaga es stilusa", "Burkolat anyaga és stílusa"),
    ("Terheles es hasznalat", "Terhelés és használat"),
    ("Vizelvezetes es szintek", "Vízelvezetés és szintek"),
    ("Elo-keszites es bontas", "Előkészítés és bontás"),
    ("Szerkezeti tipus", "Szerkezeti típus"),
    ("Fedettseg es arnyekolas", "Fedettség és árnyékolás"),
    ("Hasznalati cel", "Használati cél"),
    ("Kivant elemek", "Kívánt elemek"),
    ("Telepitesi korulmenyek", "Telepítési körülmények"),
    ("Fekves es benapozas", "Fekvés és benapozás"),
    ("Noveny karakter", "Növény karakter"),
    ("Talajelokeszites", "Talajelőkészítés"),
    ("Talajtakaras es szegely", "Talajtakarás és szegély"),
    ("Uzemi igenyek", "Üzemi igények"),
    ("Novenyanyag", "Növényanyag"),
    ("Ultetesi rend", "Ültetési rend"),
    ("Talaj es uzemeltetes", "Talaj és üzemeltetés"),
    ("Faanyag es cel", "Faanyag és cél"),
    ("Rogzites es vedelem", "Rögzítés és védelem"),
    ("Anyag es profil", "Anyag és profil"),
    ("Kapcsolodo feluletek", "Kapcsolódó felületek"),
    ("Rogzites", "Rögzítés"),
    ("Modszertan", "Módszertan"),
    ("Kiindulo allapot", "Kiinduló állapot"),
    ("Tereprendezes", "Tereprendezés"),
    ("Utokezeles", "Utókezelés"),
    ("Bontando elemek", "Bontandó elemek"),
    ("Kivitelezes modja", "Kivitelezés módja"),
    ("Anyagkezeles", "Anyagkezelés"),
    ("Helyszini adottsagok", "Helyszíni adottságok"),
    ("Szerkezeti kiegeszitok", "Szerkezeti kiegészítők"),
    ("Felhasznalas", "Felhasználás"),
    ("Vezerles es bovithetoseg", "Vezérlés és bővíthetőség"),
    ("Tisztitasi kor", "Tisztítási kör"),
    ("Megtartando elemek", "Megtartandó elemek"),
    ("Kovek es anyagok", "Kövek és anyagok"),
    ("Reteg es talaj", "Réteg és talaj"),
    ("Novenybeultetes", "Növénybeültetés"),
    ("Funkcio es latvany", "Funkció és látvány"),
    ("Tipus es anyag", "Típus és anyag"),
    ("Nyilaszarok es automatika", "Nyílászárók és automatika"),
    ("Telekhatar es bontas", "Telekhatár és bontás"),
    ("Alapozas es kiegeszitok", "Alapozás és kiegészítők"),
    ("To tipusa", "Tó típusa"),
    ("Szerkezet es burkolat", "Szerkezet és burkolat"),
    ("Fenntartas es biztonsag", "Fenntartás és biztonság"),
    ("Tervezesi csomag", "Tervezési csomag"),
    ("Forrasanyagok", "Forrásanyagok"),
    ("Tervezesi melyseg", "Tervezési mélység"),
    ("Egyuttmukodes", "Együttműködés"),
    ("Cel es stilus", "Cél és stílus"),
    ("Valtozatok", "Változatok"),
    ("Leadando anyag", "Leadandó anyag"),
    ("Datum", "Dátum"),
    ("Helyszin", "Helyszín"),
    ("helyszin", "helyszín"),
    ("helyszini", "helyszíni"),
    ("merete", "mérete"),
    ("terulet", "terület"),
    ("Terulet", "Terület"),
    ("Tervezesi", "Tervezési"),
    ("tervezesi", "tervezési"),
    ("Kiszallasi", "Kiszállási"),
    ("kiszallasi", "kiszállási"),
    ("tavolsag", "távolság"),
    ("kezdes", "kezdés"),
    ("befejezes", "befejezés"),
    ("telephelytol", "telephelytől"),
    ("telepulet", "település"),
    ("Telefonszam", "Telefonszám"),
    ("telefonszam", "telefonszám"),
    ("Kapcsolattarto", "Kapcsolattartó"),
    ("kapcsolattarto", "kapcsolattartó"),
    ("Helyszini", "Helyszíni"),
    ("helyszini", "helyszíni"),
    ("Lakoovezeti", "Lakóövezeti"),
    ("lakoovezeti", "lakóövezeti"),
    ("besorolas", "besorolás"),
    ("Bevonasa", "Bevonása"),
    ("bevonasa", "bevonása"),
    ("karbantartas", "karbantartás"),
    ("karbantartasra", "karbantartásra"),
    ("szama", "száma"),
    ("Szama", "Száma"),
    ("Kezi", "Kézi"),
    ("kezi", "kézi"),
    ("Gepi", "Gépi"),
    ("gepi", "gépi"),
    ("kivitelezes", "kivitelezés"),
    ("Kivitelezes", "Kivitelezés"),
    ("kivitelezesi", "kivitelezési"),
    ("Kivitelezesi", "Kivitelezési"),
    ("kiviteli", "kiviteli"),
    ("koltsegkeret", "költségkeret"),
    ("Koltsegkeret", "Költségkeret"),
    ("koltsegvetes", "költségvetés"),
    ("becsult", "becsült"),
    ("Becsult", "Becsült"),
    ("Gyepzonak", "Gyepzónák"),
    ("gyepzonak", "gyepzónák"),
    ("Helyszin", "Helyszín"),
    ("helyszin", "helyszín"),
    ("terfogat", "térfogat"),
    ("Terfogat", "Térfogat"),
    ("szelesseg", "szélesség"),
    ("Szelesseg", "Szélesség"),
    ("magassag", "magasság"),
    ("Magassag", "Magasság"),
    ("melyseg", "mélység"),
    ("Melyseg", "Mélység"),
    ("termofold", "termőföld"),
    ("Termofold", "Termőföld"),
    ("gyoker", "gyökér"),
    ("Gyoker", "Gyökér"),
    ("kozmu", "közmű"),
    ("Kozmu", "Közmű"),
    ("kozmukozeli", "közműközeli"),
    ("hasznalok", "használók"),
    ("Hasznalok", "Használók"),
    ("csalad", "család"),
    ("Csalad", "Család"),
    ("letszama", "létszáma"),
    ("Letszama", "Létszáma"),
    ("hazi", "házi"),
    ("Hazi", "Házi"),
    ("allat", "állat"),
    ("Allat", "Állat"),
    ("szallitasi", "szállítási"),
    ("Szallitasi", "Szállítási"),
    ("anyaglerako", "anyaglerakó"),
    ("Anyaglerako", "Anyaglerakó"),
    ("telektisztitas", "telektisztítás"),
    ("Telektisztitas", "Telektisztítás"),
    ("aramforras", "áramforrás"),
    ("Aramforras", "Áramforrás"),
    ("vilagitas", "világítás"),
    ("Vilagitas", "Világítás"),
    ("soveny", "sövény"),
    ("Soveny", "Sövény"),
    ("ultetes", "ültetés"),
    ("Ultetes", "Ültetés"),
    ("ultetesi", "ültetési"),
    ("Ultetesi", "Ültetési"),
    ("ontozes", "öntözés"),
    ("Ontozes", "Öntözés"),
    ("ontozo", "öntöző"),
    ("Ontozo", "Öntöző"),
    ("agyas", "ágyás"),
    ("Agyas", "Ágyás"),
    ("diszkavics", "díszkavics"),
    ("Diszkavics", "Díszkavics"),
    ("sziklakert", "sziklakert"),
    ("megjegyzes", "megjegyzés"),
    ("Megjegyzes", "Megjegyzés"),
    ("eszrevetel", "észrevétel"),
    ("Eszrevetel", "Észrevétel"),
    ("tervezesi", "tervezési"),
    ("Tervezesi", "Tervezési"),
    ("egyuttmukodes", "együttműködés"),
    ("Egyuttmukodes", "Együttműködés"),
    ("forrasanyagok", "forrásanyagok"),
    ("Forrasanyagok", "Forrásanyagok"),
    ("megoldasi", "megoldási"),
    ("Megoldasi", "Megoldási"),
    ("muszaki", "műszaki"),
    ("Muszaki", "Műszaki"),
    ("szint", "szint"),
    ("esztetikai", "esztétikai"),
    ("Esztetikai", "Esztétikai"),
    ("Donteshozo", "Döntéshozó"),
    ("Szallas", "Szállás"),
    ("biztositott", "biztosított"),
    ("Kapuszel.", "Kapuszél."),
    ("behajtasi", "behajtási"),
    ("nyilas", "nyílás"),
    ("Szukulet", "Szűkület"),
    ("lepcso", "lépcső"),
    ("lejto", "lejtő"),
    ("akadalyozo", "akadályozó"),
    ("Gepbejutas", "Gépbejárás"),
    ("Anyaglerako", "Anyaglerakó"),
    ("Kozmuvek", "Közművek"),
    ("rendelkezesre all", "rendelkezésre áll"),
    ("dokumentum", "dokumentum"),
    ("Geodeziai", "Geodéziai"),
    ("geodeziai", "geodéziai"),
    ("epitesi", "építési"),
    ("latvanyterv", "látványterv"),
    ("Latvanyterv", "Látványterv"),
    ("fotok", "fotók"),
    ("inspiracios", "inspirációs"),
    ("kivitelezesi", "kivitelezési"),
    ("megjegyzese", "megjegyzése"),
    ("Csalad", "Család"),
    ("letszama", "létszáma"),
    ("hasznalok", "használók"),
    ("Gyermek", "Gyermek"),
    ("Eletkor", "Életkor"),
    ("Hazi allat", "Háziállat"),
    ("sutese-fuzes", "sütés-főzés"),
    ("furdoyes", "fürdőzés"),
    ("reprezentativ", "reprezentatív"),
    ("Stilus", "Stílus"),
    ("mediterran", "mediterrán"),
    ("japan", "japán"),
    ("tropusi", "trópusi"),
    ("termeszetkozeli", "természetközeli"),
    ("kortars", "kortárs"),
    ("takaras", "takarás"),
    ("Hatsokert", "Hátsókert"),
    ("szomszed", "szomszéd"),
    ("belathatosaga", "beláthatósága"),
    ("Epulettel", "Épülettel"),
    ("eros", "erős"),
    ("igenyelt", "igényelt"),
    ("Zavaro", "Zavaró"),
    ("latvany", "látvány"),
    ("automata ontozorendszer", "automata öntözőrendszer"),
    ("robotfunyiro", "robotfűnyíró"),
    ("kertesz", "kertész"),
    ("szezonalis", "szezonális"),
    ("Koltsegkeret", "Költségkeret"),
    ("becsult", "becsült"),
    ("Megvalositas", "Megvalósítás"),
    ("onero", "önerő"),
    ("Mas", "Más"),
    ("vallalkozo", "vállalkozó"),
    ("szabalyozasi", "szabályozási"),
    ("Megrendelo", "Megrendelő"),
    ("fo", "fő"),
    ("problema", "probléma"),
    ("barmi", "bármi"),
    ("amirol", "amiről"),
    ("szo", "szó"),
    ("kockazat", "kockázat"),
    ("figyelmeztetes", "figyelmeztetés"),
    ("felmero", "felmérő"),
    ("reszerol", "részéről"),
    ("Gyepzonak", "Gyepzónák"),
    ("Agyas", "Ágyás"),
    ("zonak", "zónák"),
    ("Viznyomas", "Víznyomás"),
    ("Vizhazom", "Vízhozam"),
    ("Csovezes", "Csövezés"),
    ("Vezero", "Vezérlő"),
    ("Vizforras", "Vízforrás"),
    ("Kulon", "Külön"),
    ("gyep", "gyep"),
    ("csepegteto", "csepegtető"),
    ("mikroszorofej", "mikroszórófej"),
    ("szorofej", "szórófej"),
    ("Nyari", "Nyári"),
    ("esoerzekelo", "esőérzékelő"),
    ("Talajnedvesseg", "Talajnedvesség"),
    ("Zonankenti", "Zónánkénti"),
    ("Kabelarok", "Kábelárok"),
    ("bontas", "bontás"),
    ("atfuras", "átfúrás"),
    ("Teli", "Téli"),
    ("uzembe helyezes", "üzembe helyezés"),
    ("Kiegeszitok", "Kiegészítők"),
    ("vizoraveteli", "vízóravételi"),
    ("kulso", "külső"),
    ("szuro", "szűrő"),
    ("nyomascsokkento", "nyomáscsökkentő"),
    ("hozammeres", "hozammérés"),
    ("szerviz naplo", "szerviz napló"),
    ("Szegely", "Szegély"),
    ("vastagsag", "vastagság"),
    ("Szintkulonbseg", "Szintkülönbség"),
    ("Vizelvezetes", "Vízelvezetés"),
    ("Megkozelites", "Megközelítés"),
    ("termesko", "terméskő"),
    ("Felulethatas", "Felülethatás"),
    ("rusztikus", "rusztikus"),
    ("csuszasmentes", "csúszásmentes"),
    ("Terheles", "Terhelés"),
    ("szemelyauto", "személyautó"),
    ("kisteher", "kisteher"),
    ("fokozott", "fokozott"),
    ("Fordulo", "Forduló"),
    ("tolatasi", "tolatási"),
    ("zona", "zóna"),
    ("kitett", "kitett"),
    ("retegend", "rétegrend"),
    ("uj alapeptes", "új alapépítés"),
    ("meglevo", "meglévő"),
    ("ujra-szintezes", "újraszintezés"),
    ("Agyazohomok", "Ágyazóhomok"),
    ("geotextilia", "geotextília"),
    ("pontszeru", "pontszerű"),
    ("folyoka", "folyóka"),
    ("drencso", "dréncső"),
    ("elszikkasztas", "elszikkasztás"),
    ("Hazfal", "Házfal"),
    ("kuszob", "küszöb"),
    ("garazs", "garázs"),
    ("csatlakozas", "csatlakozás"),
    ("kritikus", "kritikus"),
    ("Elszallitas", "Elszállítás"),
    ("veszteseggel", "veszteséggel"),
    ("szamolni", "számolni"),
    ("aranyekolo", "árnyékoló"),
    ("Magassag", "Magasság"),
    ("Alapterulet", "Alapterület"),
    ("Oszlopok", "Oszlopok"),
    ("Arnyekelo", "Árnyékoló"),
    ("alu", "alu"),
    ("lamellas", "lamellás"),
    ("vaszon", "vászon"),
    ("uveg", "üveg"),
    ("novenyfuttatas", "növényfuttatás"),
    ("Esovedelem", "Esővédelem"),
    ("beepitett", "beépített"),
    ("futes", "fűtés"),
    ("oldalarnyekolo", "oldalárnyékoló"),
    ("Alapozas", "Alapozás"),
    ("vizelvezetes", "vízelvezetés"),
    ("Jatszoteri", "Játszótéri"),
    ("Felhasznalok", "Felhasználók"),
    ("Esegcsillapito", "Eséscsillapító"),
    ("Arnyekolt", "Árnyékolt"),
    ("korosztaly", "korosztály"),
    ("vendegeknek", "vendégeknek"),
    ("ovi", "óvi"),
    ("intezmeny", "intézmény"),
    ("Mozgas", "Mozgás"),
    ("egyensuly", "egyensúly"),
    ("maszas", "mászás"),
    ("kreativ", "kreatív"),
    ("csuszda", "csúszda"),
    ("maszoka", "mászóka"),
    ("hintaagy", "hintaágy"),
    ("vizijatek", "vízijáték"),
    ("haziko", "házikó"),
    ("sportos", "sportos"),
    ("Biztonsag", "Biztonság"),
    ("fakereg", "fakéreg"),
    ("Belathatosag", "Beláthatóság"),
    ("hazbol", "házból"),
    ("felszini", "felszíni"),
    ("kozmuvan", "közmű van"),
    ("Minositett", "Minősített"),
    ("szabvanyos", "szabványos"),
]

UNDERSCORE_RE = re.compile(r"_{10,}")


SECTIONS = [
    {
        "title": "Alapadatok",
        "metrics": [
            "Datum",
            "Helyszin teljes cime",
            "Telek merete (m2)",
            "Tervezesi terulet (m2)",
            "Kiszallasi tavolsag (km)",
            "Tervezett kezdes / befejezes",
        ],
        "groups": [
            {
                "title": "Kapcsolat es helyszin",
                "lines": [
                    "Kapcsolattarto neve: ________________________________________________",
                    "Telefonszam / e-mail: ________________________________________________",
                    "Helyszini kapcsolattarto: ____________________________________________",
                    "Lakoovezeti besorolas: ______________________________________________",
                    "Donteshozo: □ tulajdonos □ hazaspar □ ceg □ egyeb: ____________",
                    "Szallas biztositott: □ igen □ nem",
                ],
            },
            {
                "title": "Megkozelites es kozmuvek",
                "lines": [
                    "Kapuszel. / behajtasi nyilas: ________________________________________",
                    "Szukulet, lepcso, lejto, egyeb akadalyozo tenyezo: ____________________",
                    "Gepbejutas lehetseges: □ igen □ reszben □ nem",
                    "Anyaglerako hely kijelolheto: □ igen □ reszben □ nem",
                    "Kozmuvek ismert helye rendelkezesre all: □ igen □ reszben □ nem",
                    "Parkolas a helyszinen / kozelben megoldott: □ igen □ nem",
                ],
            },
            {
                "title": "Atadott dokumentumok",
                "lines": [
                    "□ tervrajz □ geodeziai felmeres □ helyszinrajz □ epitesi engedely",
                    "□ latvanyterv □ fotok □ inspiracios kepek □ kivitelezesi terv",
                    "□ kozmu terv □ geologia / talaj adat □ egyeb dokumentum",
                    "Dokumentumok rovid megjegyzese: _____________________________________",
                ],
            },
            {
                "title": "Hasznalati igenyek",
                "lines": [
                    "Csalad letszama / hasznalok szama: ____________________________________",
                    "Gyermek: □ nincs □ van   Eletkor: _________________________________",
                    "Hazi allat: □ nincs □ kutya □ macska □ egyeb: __________________",
                    "Celok: □ pihenes □ kozossegi ter □ sutese-fuzes □ jatszohely",
                    "□ konyhakert □ tarolas □ parkolas □ arnyekolas □ furdoyes / vizjatek",
                    "□ ures gyepfelulet □ reprezentativ latvany □ egyeb: _______________",
                ],
            },
            {
                "title": "Stilus es terkapcsolat",
                "lines": [
                    "Stilus: □ modern □ mediterran □ japan □ angol □ tropusi",
                    "□ termeszetkozeli □ minimalista □ kortars □ egyedi",
                    "Oldalso takaras szukseges: □ igen □ reszben □ nem",
                    "Hatsokert / szomszed belathatosaga zavarja: □ igen □ nem",
                    "Epulettel eros kertkapcsolat igenyelt: □ igen □ nem",
                    "Zavaro latvany eltakarasa szukseges: □ igen □ nem",
                ],
            },
            {
                "title": "Fenntartas es uzemeltetes",
                "lines": [
                    "□ automata ontozorendszer □ robotfunyiro □ sajat fenntartas",
                    "□ rendszeres kertesz □ alkalmi kertesz □ szezonalis karbantartas",
                    "Fenntartasi szint: □ alacsony □ kozepes □ magas",
                    "Uzemeltetesi megjegyzes: _____________________________________________",
                ],
            },
            {
                "title": "Koltsegkeret es utemezes",
                "lines": [
                    "Koltsegkeret: □ 1 M alatt □ 1-2 M □ 2-4 M □ 4-7 M □ 7-10 M □ 10 M+",
                    "Pontos koltsegkeret / becsult plafon: _________________________________",
                    "Megvalositas: □ teljes kivitelezes □ szakaszos □ onero + szakipar",
                    "Mas vallalkozoval is dolgozik: □ igen □ nem",
                    "HESZ / szabalyozasi ellenorzes szukseges: □ igen □ nem",
                ],
            },
            {
                "title": "Projektosszefoglalo",
                "lines": [
                    "Megrendelo fo celja: _________________________________________________",
                    "Legfontosabb problema: _______________________________________________",
                    "Van-e barmi, amirol meg nem esett szo: ______________________________",
                    "Kiemelt kockazat / figyelmeztetes a felmero reszerol: ________________",
                ],
            },
        ],
    },
    {
        "title": "Ontozorendszer",
        "metrics": [
            "Gyepzonak szama",
            "Agyas zonak szama",
            "Viznyomas (bar)",
            "Vizhazom / hozam (l/perc)",
            "Csovezes hossza (fm)",
            "Vezero korok szama",
        ],
        "groups": [
            {
                "title": "Rendszer jellege",
                "lines": [
                    "Tipus: □ uj rendszer □ bovites □ felujitas □ javitas □ szerviz",
                    "Vizforras: □ halozati □ kut □ eso-viz tarolo □ egyeb",
                    "Kulon gyep es agyas korok: □ igen □ reszben □ nem",
                    "Jelenlegi rendszer a helyszinen: □ nincs □ van □ reszben mukodik",
                ],
            },
            {
                "title": "Ontozesi mod",
                "lines": [
                    "□ rotoros szorofej □ spray szorofej □ csepegteto □ mikroszorofej",
                    "□ cserepes ontozes □ fakor ontozes □ noveny-agyas ontozes",
                    "Nyari csucsigeny / intenziv ontozes szukseges: □ igen □ nem",
                ],
            },
            {
                "title": "Automatizalas",
                "lines": [
                    "Vezerles: □ manualis □ automata □ wifi / okosvezerles",
                    "Esoerzekelo: □ igen □ nem   Talajnedvesseg erzekelo: □ igen □ nem",
                    "Zonankenti kulon idozites igenyelt: □ igen □ nem",
                ],
            },
            {
                "title": "Kivitelezesi korulmenyek",
                "lines": [
                    "Kabelarok / csatorna egyeb kozmu mellett: □ igen □ nem",
                    "Burkolatbontas vagy utolagos atfuras szukseges: □ igen □ nem",
                    "Teli vizesites / leeresztes / tavaszi uzembe helyezes kell: □ igen □ nem",
                ],
            },
            {
                "title": "Kiegeszitok",
                "lines": [
                    "□ vizoraveteli pont □ kulso csap □ akna □ szuro □ nyomascsokkento",
                    "□ kut es halozat kozti valtas □ hozammeres □ szerviz naplo",
                ],
            },
        ],
    },
    {
        "title": "Burkolatok",
        "metrics": [
            "Burkolando felulet (m2)",
            "Szegely hossza (fm)",
            "Burkolat vastagsag (cm)",
            "Szintkulonbseg (cm)",
            "Vizelvezetes hossza (fm)",
            "Megkozelites: kezi / gepi",
        ],
        "groups": [
            {
                "title": "Burkolat anyaga es stilusa",
                "lines": [
                    "□ viacolor □ nagylap □ termeskő □ betonlap □ fa / WPC □ murva",
                    "Felhasznalas: □ jarda □ terasz □ kocsibeallo □ lepcso □ piheno",
                    "Felulethatas: □ natur □ modern □ rusztikus □ csuszasmentes",
                ],
            },
            {
                "title": "Terheles es hasznalat",
                "lines": [
                    "Terheles: □ gyalogos □ szemelyauto □ kisteher □ fokozott terheles",
                    "Fordulo / tolatasi zona van: □ igen □ nem",
                    "Fagy- es csapadekterheles erosen kitett: □ igen □ nem",
                ],
            },
            {
                "title": "Alapozas es retegend",
                "lines": [
                    "□ uj alapeptes □ meglevo alapra □ javitas □ ujra-szintezes",
                    "Alapretegek: □ zuzottko □ beton □ agyazohomok □ ragasztott rendszer",
                    "Geotextilia szukseges: □ igen □ nem",
                ],
            },
            {
                "title": "Vizelvezetes es szintek",
                "lines": [
                    "Lejtes ellenorzese: □ igen □ nem",
                    "□ pontszeru vizelvezetes □ folyoka □ drencso □ elszikkasztas",
                    "Hazfalhoz, kuszobhoz vagy garazshoz csatlakozas kritikus: □ igen □ nem",
                ],
            },
            {
                "title": "Elo-keszites es bontas",
                "lines": [
                    "Meglevo burkolat bontasa: □ igen □ nem",
                    "Elszallitas: □ igen □ nem   Vagasi veszteseggel kell szamolni: □ igen □ nem",
                    "Szegely tipus: □ beton □ acel □ muanyag □ termeskő □ egyeb",
                ],
            },
        ],
    },
    {
        "title": "Pergolak",
        "metrics": [
            "Hossz (m)",
            "Szelesseg (m)",
            "Magassag (m)",
            "Alapterulet (m2)",
            "Oszlopok szama (db)",
            "Arnyekolo mezok (db)",
        ],
        "groups": [
            {
                "title": "Szerkezeti tipus",
                "lines": [
                    "□ szabadon allo □ hazhoz kotott □ falra rogzitett □ egyedi acelszerkezet",
                    "Anyag: □ fa □ fem □ alu □ WPC □ vegyes",
                    "Kialakitas: □ modern □ klasszikus □ lamellas □ textil arnyekolos",
                ],
            },
            {
                "title": "Fedettseg es arnyekolas",
                "lines": [
                    "Tetozet: □ nyitott □ policarbonat □ uveg □ zsaluzias □ vaszon",
                    "Oldalfal / takaras: □ nincs □ zsaluzias □ uveg □ textil □ novenyfuttatas",
                    "Esovedelem teljes erteku legyen: □ igen □ nem",
                ],
            },
            {
                "title": "Kiegeszitok",
                "lines": [
                    "□ beepitett vilagitas □ konnektor □ futes □ ventilator □ hangtechnika",
                    "□ oldalarnyekolo □ zip screen □ novenyfuttato huzalok",
                ],
            },
            {
                "title": "Alapozas es csatlakozas",
                "lines": [
                    "Alapozas: □ pont alap □ savalap □ burkolt feluletre rogzitett",
                    "Vizelvezetes megoldasa szukseges: □ igen □ nem",
                    "Csatlakozas meglevo teraszhoz / burkolathoz: □ igen □ nem",
                ],
            },
        ],
    },
    {
        "title": "Jatszoteri elemek",
        "metrics": [
            "Terulet (m2)",
            "Felhasznalok korcsoportja",
            "Elemek szama (db)",
            "Esegcsillapito felulet (m2)",
            "Biztonsagi zona (m2)",
            "Arnyekolt felulet (m2)",
        ],
        "groups": [
            {
                "title": "Hasznalati cel",
                "lines": [
                    "Korcsoport: □ 0-3 □ 3-6 □ 6-12 □ tobb korosztaly",
                    "Hasznalat: □ csaladi □ kozossegi □ vendegeknek □ ovi / intezmeny",
                    "Prioritas: □ mozgas □ egyensuly □ maszas □ kreativ jatek",
                ],
            },
            {
                "title": "Kivant elemek",
                "lines": [
                    "□ hinta □ csuszda □ maszoka □ homokozo □ trambulin □ hintaagy",
                    "□ vizijatek □ rugos jatek □ kis hazikó □ sportos elem",
                ],
            },
            {
                "title": "Biztonsag",
                "lines": [
                    "Esegesillapito reteg: □ gumilap □ fakereg □ homok □ fuvesitett □ egyeb",
                    "Korbe kerites szukseges: □ igen □ nem",
                    "Belathatosag a hazbol fontos: □ igen □ nem",
                ],
            },
            {
                "title": "Telepitesi korulmenyek",
                "lines": [
                    "Felszini akadaly / gyoker / kozmu van: □ igen □ nem",
                    "Arnyekolas kell: □ igen □ reszben □ nem",
                    "Minositett / szabvanyos rendszer szukseges: □ igen □ nem",
                ],
            },
        ],
    },
    {
        "title": "Agyasok",
        "metrics": [
            "Agyas hossza (fm)",
            "Agyas szelessege (m)",
            "Osszterulet (m2)",
            "Talajcsere (m3)",
            "Szegely hossza (fm)",
            "Ontozott terulet (m2)",
        ],
        "groups": [
            {
                "title": "Fekves es benapozas",
                "lines": [
                    "□ napos □ felarnyekos □ arnyekos □ szeles □ vedett",
                    "Agyas funkcioja: □ disz □ takaro □ eves erdeklodes □ bejarati hangsuly",
                    "Meglevo novenyzet megtartando: □ igen □ reszben □ nem",
                ],
            },
            {
                "title": "Noveny karakter",
                "lines": [
                    "□ homogen □ tobbfajta □ alacsony □ kozepes □ magas □ viragzo □ orokzold",
                    "Kedvelt novenyek / tiltott novenyek: ___________________________________",
                    "Illat, szezonhatas, szinpaletta fontos: □ igen □ reszben □ nem",
                ],
            },
            {
                "title": "Talajelokeszites",
                "lines": [
                    "□ talajcsere □ lazitas □ komposzt □ marhatragya □ savanyitas □ drenezes",
                    "Gyommentesites: □ kezi □ geotextil □ karton □ vegyszermentes □ vegyszeres",
                ],
            },
            {
                "title": "Talajtakaras es szegely",
                "lines": [
                    "□ fakereg □ faapritek □ díszkavics □ murva □ kulé □ novenytakaro",
                    "Szegely: □ betonba rakott □ szarazon rakott □ acel □ muanyag □ termeskő",
                ],
            },
            {
                "title": "Uzemi igenyek",
                "lines": [
                    "Ontozes: □ nincs □ csepegteto □ mikroszorofej □ kezis ontozes",
                    "Fenntartasi igeny: □ alacsony □ kozepes □ intenziv",
                ],
            },
        ],
    },
    {
        "title": "Sovenyek",
        "metrics": [
            "Hossz (fm)",
            "Sorok szama",
            "Totav (cm)",
            "Novenyek szama (db)",
            "Vegleges magassag (cm)",
            "Ontozott hossza (fm)",
        ],
        "groups": [
            {
                "title": "Funkcio",
                "lines": [
                    "□ belatasgatlas □ telekhatar □ szelvedelem □ dekoracios □ zajcsokkento",
                    "Suru, gyors takaras igenyelt: □ igen □ reszben □ nem",
                    "Egesz eves takaras kell: □ igen □ nem",
                ],
            },
            {
                "title": "Novenyanyag",
                "lines": [
                    "□ lombhullato □ orokzold □ nyirott □ termeszetes hatasu",
                    "Faj / preferencia: _________________________________________________",
                    "Meglevo sovennyel egyutt kezelendo: □ igen □ nem",
                ],
            },
            {
                "title": "Ultetesi rend",
                "lines": [
                    "□ egysoros □ ketsoros □ haromszogkotés □ valtozo ritmus",
                    "Magassag mar uleskor: □ kicsi □ kozepes □ nagymeretu anyag",
                    "Gyokertereles vagy kozmuvedelem szukseges: □ igen □ nem",
                ],
            },
            {
                "title": "Talaj es uzemeltetes",
                "lines": [
                    "□ talajcsere □ komposzt □ marhatragya □ geotextil □ mulcs",
                    "Ontozes: □ nincs □ csepegteto □ kezis □ automata korbe kotve",
                    "Fenntartasi elvaras: □ ritka nyiras □ formalis nyiras □ gyors novekedes",
                ],
            },
        ],
    },
    {
        "title": "Fa ultetes",
        "metrics": [
            "Fak szama (db)",
            "Torzskormeret (cm)",
            "Kontener / labda meret",
            "Ultetogodor (db)",
            "Talajcsere (m3)",
            "Kikotok szama (db)",
        ],
        "groups": [
            {
                "title": "Faanyag es cel",
                "lines": [
                    "Tipus: □ diszfa □ arnyekfa □ gyumolcsfa □ fasor □ hangsulyos szoliter",
                    "Lombhullato / orokzold: □ lombhullato □ orokzold",
                    "Cel: □ arnyek □ latvany □ takaras □ karakteres bejarat □ termes",
                ],
            },
            {
                "title": "Elhelyezes",
                "lines": [
                    "Epulettol / kozmutol tavolsag kritikus: □ igen □ nem",
                    "Burkolat kozeleben ultetve: □ igen □ nem",
                    "Fak koronaja utkozik: □ vezetek □ tetosik □ szomszed □ nem",
                ],
            },
            {
                "title": "Talaj es ultetes",
                "lines": [
                    "□ talajcsere □ komposzt □ marhatragya □ gyokerterelo □ drenezes",
                    "□ ontozotanyer □ ontozozsak □ fakor kialakitas □ takarokereg",
                ],
            },
            {
                "title": "Rogzites es vedelem",
                "lines": [
                    "□ kikoto karok □ heveder □ torzsvedelem □ nyul / kutya elleni vedelem",
                    "Kezdo ontozes / beiszapolas biztositas: □ igen □ nem",
                ],
            },
        ],
    },
    {
        "title": "Szegely epites",
        "metrics": [
            "Szegely hossza (fm)",
            "Ives szakasz (fm)",
            "Sarokpontok (db)",
            "Betonagy (m3)",
            "Csatlakozo agyas / burkolat (db)",
            "Bontas szukseges-e",
        ],
        "groups": [
            {
                "title": "Anyag es profil",
                "lines": [
                    "□ beton szegely □ acel □ muanyag □ termeskő □ tegla □ WPC",
                    "Kialakitas: □ egyenes □ ivelt □ rejtett □ magas szegely",
                ],
            },
            {
                "title": "Kapcsolodo feluletek",
                "lines": [
                    "□ burkolat mellett □ gyep mellett □ agyas mellett □ murva mellett",
                    "Szin / anyag harmonizacio fontos: □ igen □ nem",
                ],
            },
            {
                "title": "Rogzites",
                "lines": [
                    "□ betonba rakott □ szarazon rakott □ cövekelt □ ragasztott",
                    "Geotextil / gyomvedelem csatlakozik: □ igen □ nem",
                ],
            },
            {
                "title": "Kivitelezesi korlatok",
                "lines": [
                    "Gyoker, kozmu vagy egyeb akadalyozo tenyezo: □ igen □ nem",
                    "Utolagos szintezes / kiegyenlites szukseges: □ igen □ nem",
                ],
            },
        ],
    },
    {
        "title": "Fuvesites",
        "metrics": [
            "Terulet (m2)",
            "Gyep szel hossza (fm)",
            "Termofold (m3)",
            "Komposzt (m3)",
            "Szintkulonbseg (cm)",
            "Ontozes elo-keszites",
        ],
        "groups": [
            {
                "title": "Modszertan",
                "lines": [
                    "□ fumagvetes □ gyepszonyeg □ javito felulvetes □ reszleges potlas",
                    "Gyors hasznalhatosag fontos: □ igen □ nem",
                ],
            },
            {
                "title": "Kiindulo allapot",
                "lines": [
                    "□ gyomos □ kopar □ toredezett □ meglevo gyep javitando □ geotextiles",
                    "Aljnovenyzet eltavolitasa kell: □ igen □ nem",
                ],
            },
            {
                "title": "Talajelokeszites",
                "lines": [
                    "□ gyomirtas □ felso-maras □ rotalas □ komposzt □ marhatragya",
                    "□ talajcsere □ fold atcsoportositas □ elszallitas □ termofold terites",
                ],
            },
            {
                "title": "Tereprendezes",
                "lines": [
                    "□ durva tereprendezes geppel □ durva tereprendezes kezzel",
                    "□ finom tereprendezes □ szintezo vassal □ vizmertekes ellenorzes",
                    "Vizelvezetes kulon figyelmet igenyel: □ igen □ nem",
                ],
            },
            {
                "title": "Utokezeles",
                "lines": [
                    "□ ontozes programozasa □ starter tap □ gyepgondozasi utmutato",
                    "□ utolagos hengerles □ robotfunyiro elo-keszites",
                ],
            },
        ],
    },
    {
        "title": "Bontas",
        "metrics": [
            "Bontando felulet (m2)",
            "Vastagsag (cm)",
            "Sitt mennyiseg (m3)",
            "Rakodasi tav (m)",
            "Lepcsok / szintek (db)",
            "Elszallitasi fordulok",
        ],
        "groups": [
            {
                "title": "Bontando elemek",
                "lines": [
                    "□ burkolat □ beton □ tegla □ tamfal □ kerites □ pergola □ tarolo",
                    "□ tusko □ fa □ gyokerzet □ kozmukozeli szakasz □ egyeb",
                ],
            },
            {
                "title": "Kivitelezes modja",
                "lines": [
                    "□ kezi bontas □ gepi bontas □ vagassal □ fureszelessel □ darabolassal",
                    "Zaj / ido korlatozas van: □ igen □ nem",
                ],
            },
            {
                "title": "Anyagkezeles",
                "lines": [
                    "□ szelektiv szetvalogatas □ ujrahasznositas □ kontener □ helyszini rakodas",
                    "Elszallitas megrendelve: □ igen □ nem",
                ],
            },
            {
                "title": "Kockazatok",
                "lines": [
                    "Kozmuves / gyokeres / epuletszegely melletti bontas: □ igen □ nem",
                    "Megtartando elemek vedelme kell: □ igen □ nem",
                ],
            },
        ],
    },
    {
        "title": "Tamfal epites",
        "metrics": [
            "Hossz (fm)",
            "Atlagmagassag (m)",
            "Max. magassag (m)",
            "Foldmunka (m3)",
            "Drenhossz (fm)",
            "Lepcsok / torések (db)",
        ],
        "groups": [
            {
                "title": "Rendszer",
                "lines": [
                    "□ zsalukő □ monolit beton □ termeskő □ gabion □ novenykazettas",
                    "□ szarazrakott ko □ feszitett rendszer □ egyedi statikai szerkezet",
                ],
            },
            {
                "title": "Helyszini adottsagok",
                "lines": [
                    "Retegzett, omlos vagy nedves talaj: □ igen □ nem",
                    "Jelentos szintkulonbseg megtartasa a cel: □ igen □ nem",
                    "Szomszed / epulet kozel van: □ igen □ nem",
                ],
            },
            {
                "title": "Szerkezeti kiegeszitok",
                "lines": [
                    "□ drenezes □ hatso visszatoltes □ georacs □ vizkivezetes □ korlat",
                    "□ lepcso integralas □ novenykazetta □ fedko □ burkolt csatlakozas",
                ],
            },
            {
                "title": "Alapozas es kivitelezes",
                "lines": [
                    "□ alapozas szukseges □ megl. alapra □ fokozott gepepitmeny kell",
                    "Statikai tervezes igenyelt: □ igen □ nem",
                ],
            },
        ],
    },
    {
        "title": "Kerti vilagitas",
        "metrics": [
            "Lampatestek (db)",
            "Kapcsolasi korok (db)",
            "Kabelhossz (fm)",
            "Transzformatorok (db)",
            "Felszallasok (db)",
            "Talajba sullyesztett pontok (db)",
        ],
        "groups": [
            {
                "title": "Vilagitasi cel",
                "lines": [
                    "□ diszfeny □ kozlekedesi feny □ hangulatfény □ biztonsagi feny",
                    "□ fak / novenyek kiemelese □ burkolat jelolese □ tamfal hangsuly",
                ],
            },
            {
                "title": "Lampatest tipusok",
                "lines": [
                    "□ allolampa □ leszuro spot □ fali □ talajba sullyesztett □ szalag LED",
                    "□ lepcsofeny □ oszlopfény □ vizalatti világitas",
                ],
            },
            {
                "title": "Vezerles",
                "lines": [
                    "□ kezi kapcsolas □ idozito □ fenyerzekelo □ okosotthon □ tavvezerles",
                    "Kulon jelenetek / zonak kellenek: □ igen □ nem",
                ],
            },
            {
                "title": "Kivitelezesi korulmenyek",
                "lines": [
                    "Burkolat / gyep / agyas alatti vezetes: □ igen □ nem",
                    "Meglevo aramforrassal osszekotheto: □ igen □ nem",
                    "Vizvedett szerelvenyek kellenek: □ igen □ nem",
                ],
            },
        ],
    },
    {
        "title": "Kerti aramforras",
        "metrics": [
            "Kiallasok szama (db)",
            "Dugaljak (db)",
            "Kabelhossz (fm)",
            "Kertei fogyasztok (db)",
            "Kulon korok (db)",
            "Terheles / teljesitmeny",
        ],
        "groups": [
            {
                "title": "Felhasznalas",
                "lines": [
                    "□ ontozo automatika □ vilagitas □ robotfunyiro □ pergola □ szivattyu",
                    "□ kerti konyha □ jacuzzi □ kapunyitas □ medence gepeszet □ toltopont",
                ],
            },
            {
                "title": "Szerelvenyek",
                "lines": [
                    "□ kulteri dugalj □ foldbe rejett kiallas □ falon kivuli □ elosztoszekreny",
                    "IP vedettsegi igeny fokozott: □ igen □ nem",
                ],
            },
            {
                "title": "Vezerles es bovithetoseg",
                "lines": [
                    "□ kulon kapcsolasi korok □ tavvezerles □ idozites □ okosvezerles",
                    "Kesobbi bovitesre elo-keszites kell: □ igen □ nem",
                ],
            },
            {
                "title": "Kivitelezesi korulmenyek",
                "lines": [
                    "Arokasasi tavolsag / burkolat atvezetes: ______________________________",
                    "Meglevo halozat elegendo: □ igen □ reszben □ nem",
                ],
            },
        ],
    },
    {
        "title": "Telektisztitas",
        "metrics": [
            "Terulet (m2)",
            "Bozotos resz (m2)",
            "Kivagando fak (db)",
            "Tuskok (db)",
            "Zoldhulladek (m3)",
            "Elszallitas fordulok",
        ],
        "groups": [
            {
                "title": "Tisztitasi kor",
                "lines": [
                    "□ bozotirtas □ gazvagas □ fakivagas □ gallyazas □ tuskozas □ gyokerkiszedes",
                    "□ invaziv faj visszaszoritas □ epitesi hulladek eltavolitas",
                ],
            },
            {
                "title": "Megtartando elemek",
                "lines": [
                    "Megtartando noveny / fa / epitett elem: _______________________________",
                    "Vedendő zonak kijelolve: □ igen □ nem",
                ],
            },
            {
                "title": "Kivitelezesi mod",
                "lines": [
                    "□ kezi □ gepi □ emelos kosaras □ daralos □ tuskomaro",
                    "Belso udvar / nehezen megkozelitheto terulet: □ igen □ nem",
                ],
            },
            {
                "title": "Anyagkezeles",
                "lines": [
                    "□ helyszini apritas □ elszallitas □ daralas utan ujrahasznositas",
                    "□ gyoker / tusko kulon kezelendo □ szelektiv hulladekkezeles",
                ],
            },
        ],
    },
    {
        "title": "Tereprendezes",
        "metrics": [
            "Hossz (m)",
            "Szelesseg (m)",
            "Terulet (m2)",
            "Szintkulonbseg (cm)",
            "Tobbletfold (m3)",
            "Foldhiany (m3)",
        ],
        "groups": [
            {
                "title": "Terulet jellemzese",
                "lines": [
                    "□ sik □ lejtős □ tagolt □ tobb szintu □ vizallasos □ tomorodott",
                    "Talaj jellemzese: □ kotott □ homokos □ agyagos □ kovicsos □ vegyes",
                ],
            },
            {
                "title": "Foldmozgas",
                "lines": [
                    "□ helyszini atcsoportositas □ elszallitas □ potlas □ depózas",
                    "Teljesen geppel megoldhato: □ igen □ reszben □ nem",
                    "Reszben kezi hordas szukseges: □ igen □ nem",
                ],
            },
            {
                "title": "Logisztika",
                "lines": [
                    "Földhordas tavolsaga geppel: __________________________________________",
                    "Földhordas tavolsaga talicskaval: _____________________________________",
                    "Vodorrel / lepcson mozgatando: □ igen □ nem",
                ],
            },
            {
                "title": "Talajjavitas",
                "lines": [
                    "□ termofold □ komposzt □ marhatragya □ talajlazitas □ drenezes",
                    "□ humuszpotlas □ szikla / tormelekszedes □ gyommentesites",
                ],
            },
            {
                "title": "Viz es szintek",
                "lines": [
                    "□ feluleti vizelvezetes □ pontszeru vizgyujtes □ szikkaszto □ drencso",
                    "Finom szintezes kulon pontossagot igenyel: □ igen □ nem",
                ],
            },
        ],
    },
    {
        "title": "Sziklakert epites",
        "metrics": [
            "Terulet (m2)",
            "Szintvaltasi magassag (cm)",
            "Kovanyag mennyiseg",
            "Novenyek (db)",
            "Szegely hossza (fm)",
            "Drenreteg (m3)",
        ],
        "groups": [
            {
                "title": "Koncepcio",
                "lines": [
                    "Stilus: □ termeszetes □ modern □ hegyvideki □ minimal □ gyujtemenyes",
                    "Funkcio: □ dekoracios □ szintkezeles □ kis fenntartas □ hangsulyos resz",
                ],
            },
            {
                "title": "Kovek es anyagok",
                "lines": [
                    "□ termeskő □ andezit □ meszko □ bazalt □ gneisz □ vegyes",
                    "Kovek merete: □ apro □ kozepes □ nagy □ sziklatomb",
                ],
            },
            {
                "title": "Reteg es talaj",
                "lines": [
                    "□ drenezes □ geotextil □ kavics alap □ sovany talaj □ novenyfold javitas",
                    "Vizelvezetes kritikus: □ igen □ nem",
                ],
            },
            {
                "title": "Novenybeultetes",
                "lines": [
                    "□ szarazsagturo □ alacsony parnazos □ pozsgas □ díszfű □ talajtakaro",
                    "Napfényviszony: □ napos □ felarnyekos □ arnyekos",
                ],
            },
        ],
    },
    {
        "title": "Diszkavics agyas epites",
        "metrics": [
            "Terulet (m2)",
            "Kavicsvastagsag (cm)",
            "Kavicsszukseglet",
            "Szegely hossza (fm)",
            "Geotextil (m2)",
            "Novenyek (db)",
        ],
        "groups": [
            {
                "title": "Funkcio es latvany",
                "lines": [
                    "□ dekoracios □ alacsony fenntartasu □ takaro □ reprezentativ",
                    "□ modern □ termeszetes □ mediterran □ japanszeru",
                ],
            },
            {
                "title": "Kovanyag",
                "lines": [
                    "□ diszkavics □ murva □ kulé □ zúzottko □ folyami kavics",
                    "Szemcsemeret: □ finom □ kozepes □ durva □ vegyes",
                    "Szin: □ feher □ szurke □ antracit □ barna □ vegyes",
                ],
            },
            {
                "title": "Alapreteg",
                "lines": [
                    "□ geotextil □ alapkiegyenlites □ talajcsere □ gyommentesites □ drenezes",
                    "Meglevo novenyek koze kerul: □ igen □ nem",
                ],
            },
            {
                "title": "Szegely es fenntartas",
                "lines": [
                    "Szegely: □ acel □ muanyag □ beton □ termeskő □ rejtett",
                    "Fenntartasi szint: □ alacsony □ kozepes □ gyomlálást igenyel",
                ],
            },
        ],
    },
    {
        "title": "Kerites epites",
        "metrics": [
            "Hossz (fm)",
            "Magassag (cm)",
            "Oszlopok (db)",
            "Kapu (db)",
            "Kiskapu (db)",
            "Alapozas hossza (fm)",
        ],
        "groups": [
            {
                "title": "Tipus es anyag",
                "lines": [
                    "□ drotfonat □ tabla □ fa □ WPC □ fem □ zsalukő / bazisfal □ gabion",
                    "Hatasa: □ teljes takaras □ atlathato □ disz □ biztonsagi",
                ],
            },
            {
                "title": "Nyilaszarok es automatika",
                "lines": [
                    "□ tolokapu □ nyilokapu □ kiskapu □ automata nyitas □ kaputelefon",
                    "Aramellatas biztositando: □ igen □ nem",
                ],
            },
            {
                "title": "Telekhatar es bontas",
                "lines": [
                    "Telekhatar egyertelmu: □ igen □ reszben □ nem",
                    "Meglevo kerites bontasa: □ igen □ nem",
                    "Szintkulonbseg miatt lepcsozes / torés kell: □ igen □ nem",
                ],
            },
            {
                "title": "Alapozas es kiegeszitok",
                "lines": [
                    "□ pontalap □ savalap □ labazat □ pillerek □ lamellak □ belatasgátlo betet",
                    "Soveny vagy novenyfuttatas kapcsolodik hozza: □ igen □ nem",
                ],
            },
        ],
    },
    {
        "title": "Kerti to epites",
        "metrics": [
            "Hossz (m)",
            "Szelesseg (m)",
            "Max. melyseg (m)",
            "Vizterfogat (m3)",
            "Gepeszeti pontok (db)",
            "Patak / vizjatek hossza (fm)",
        ],
        "groups": [
            {
                "title": "To tipusa",
                "lines": [
                    "□ diszto □ halasto □ termeszetes hatasu □ formalis □ csobogo / vizjatek",
                    "□ patak / vizestés □ tavacska □ tukorviz",
                ],
            },
            {
                "title": "Szerkezet es burkolat",
                "lines": [
                    "□ folia □ merev medence □ beton □ termosko szegely □ fa fedlap",
                    "Partkepzes: □ lapos □ lepcsozetes □ gyermekbarat □ meredek hangsuly",
                ],
            },
            {
                "title": "Gepeszet",
                "lines": [
                    "□ szuro □ szivattyu □ UV □ automata utantoltes □ levegozteto",
                    "□ vizforgatas □ teliesites □ tulfolyo / biztonsagi leeresztes",
                ],
            },
            {
                "title": "Fenntartas es biztonsag",
                "lines": [
                    "□ halas uzem □ novenyes uzem □ keves fenntartas □ intenziv latvany",
                    "Gyermekbiztonsag kiemelten fontos: □ igen □ nem",
                    "Vilagitas / aram ellatas szukseges: □ igen □ nem",
                ],
            },
        ],
    },
    {
        "title": "Tervrajz keszites",
        "metrics": [
            "Tervezesi terulet (m2)",
            "Tervezesi valtozatok (db)",
            "Modositasi korok",
            "Helyszini alkalmak",
            "Reszlettervek (db)",
            "Leadando csomag szint",
        ],
        "groups": [
            {
                "title": "Tervezesi csomag",
                "lines": [
                    "□ koncepcio □ alaprajz □ anyagjeloles □ novenykiiras □ muszaki reszletek",
                    "□ kivitelezesi csomag □ telepitési terv □ ontozesi terv □ vilagitasi terv",
                ],
            },
            {
                "title": "Forrasanyagok",
                "lines": [
                    "□ geodezia □ helyszinrajz □ epitesz terv □ kozmu terv □ fotok □ inspiracio",
                    "Meretek felveve a helyszinen: □ igen □ reszben □ nem",
                ],
            },
            {
                "title": "Tervezesi melyseg",
                "lines": [
                    "□ vazlatterv □ engedelyhez kozeli □ kivitelezesre alkalmas □ reszletes kotet",
                    "HESZ / telepuleskep figyelembevetel kell: □ igen □ nem",
                ],
            },
            {
                "title": "Egyuttmukodes",
                "lines": [
                    "□ online egyeztetes □ szemelyes egyeztetes □ helyszini bejaras □ workshop",
                    "Mas szakaggal osszehangolando: □ epitesz □ gepsz □ villany □ egyeb",
                ],
            },
        ],
    },
    {
        "title": "Latvanyterv keszites",
        "metrics": [
            "Nezetek szama (db)",
            "Kepek szama (db)",
            "Alternativak (db)",
            "Nappali / esti verziok",
            "Animacio igeny",
            "Atadasi formatum",
        ],
        "groups": [
            {
                "title": "Cel es stilus",
                "lines": [
                    "□ dontes-elokeszites □ ertekesites □ kivitelezoi tamogatas □ marketing",
                    "Stilus: □ realisztikus □ hangulati □ minimal □ fotorealisztikus",
                ],
            },
            {
                "title": "Nezetek",
                "lines": [
                    "□ utca felol □ terasz felol □ felso nezet □ bejarati fokusz □ esti kep",
                    "□ dron jellegu □ reszletkep □ novenykiemelo kep",
                ],
            },
            {
                "title": "Valtozatok",
                "lines": [
                    "□ egy koncepcio □ ket alternativ koncepcio □ anyagvarians □ novenyvarians",
                    "Modositasi korok szama: _____________________________________________",
                ],
            },
            {
                "title": "Leadando anyag",
                "lines": [
                    "□ JPG □ PDF □ prezentalhato tablak □ nyomtathato lapok □ animacio",
                    "Brandelt prezentacios igeny: □ igen □ nem",
                ],
            },
        ],
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D8CCBA", size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = tc_mar.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def style_run(run, size, bold=False, color=None):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_text(paragraph, text, size=9, bold=False, color=None):
    run = paragraph.add_run(fix_hu(text))
    style_run(run, size=size, bold=bold, color=color)
    return run


def fix_hu(text):
    fixed = text
    for source, target in HU_REPLACEMENTS:
        if " " in source or "/" in source or ":" in source or "+" in source or "-" in source:
            fixed = fixed.replace(source, target)
        else:
            fixed = re.sub(rf"(?<!\w){re.escape(source)}(?!\w)", target, fixed)
    return fixed


def paper_line(text):
    if "____" not in text:
        return text
    return UNDERSCORE_RE.sub("__________________", text, count=1)


def add_header(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    add_text(paragraph, "DISZKERTEK | BOVITETT KERTEPITOI HELYSZINI FELMERO", size=9, bold=True, color=GREEN_DARK)

    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:color"), SAND_DARK)
    border.append(bottom)
    paragraph._p.get_or_add_pPr().append(border)


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)
    add_header(section)

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    style.font.size = Pt(9)


def add_title_block(doc, title):
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_table.autofit = False
    cell = title_table.cell(0, 0)
    cell.width = Cm(25.8)
    set_cell_shading(cell, GREEN)
    set_cell_border(cell, color=GREEN, size="4")
    set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    add_text(paragraph, title, size=17, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    add_text(paragraph, "Nyomtathato, pipalhato, rovid megjegyzesekkel kitoltheto oldal.", size=8, color=TAUPE)


def add_metrics_table(doc, metrics):
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(8.55)
            set_cell_shading(cell, SAND)
            set_cell_border(cell)
            set_cell_margins(cell, top=65, start=85, bottom=65, end=85)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for idx, metric in enumerate(metrics[:6]):
        cell = table.cell(idx // 3, idx % 3)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(1)
        add_text(paragraph, metric, size=8, bold=True, color=GREEN_DARK)
        value = cell.add_paragraph()
        value.paragraph_format.space_before = Pt(0)
        add_text(value, "__________________________", size=9, color=TAUPE)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)


def add_groups_table(doc, groups):
    rows = (len(groups) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    index = 0
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(12.85)
            set_cell_shading(cell, "FBF9F5")
            set_cell_border(cell)
            set_cell_margins(cell, top=70, start=90, bottom=70, end=90)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if index >= len(groups):
                cell.text = ""
                index += 1
                continue
            group = groups[index]
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            add_text(cell.paragraphs[0], group["title"], size=9, bold=True, color=GREEN_DARK)
            for line in group["lines"]:
                paragraph = cell.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0.8)
                paragraph.paragraph_format.line_spacing = 1.0
                add_text(paragraph, paper_line(line), size=8.2, color=GRAY)
            index += 1
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)


def add_notes_box(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(25.8)
    set_cell_shading(cell, SAND)
    set_cell_border(cell)
    set_cell_margins(cell, top=80, start=100, bottom=80, end=100)
    title = cell.paragraphs[0]
    title.paragraph_format.space_after = Pt(2)
    add_text(title, "Megjegyzes / helyszini eszrevetel", size=9, bold=True, color=GREEN_DARK)
    for _ in range(3):
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        add_text(paragraph, "____________________________________________________________________________________", size=8.5, color=TAUPE)


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Diszkertek bovitett helyszini felmero"
    doc.core_properties.subject = "Kertepitoi helyszini felmero lap"
    doc.core_properties.author = "Codex"

    for idx, section_data in enumerate(SECTIONS):
        if idx > 0:
            doc.add_page_break()
        add_title_block(doc, section_data["title"])
        add_metrics_table(doc, section_data["metrics"])
        add_groups_table(doc, section_data["groups"])
        add_notes_box(doc)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
